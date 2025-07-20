import tkinter as tk
from tkinter import filedialog
import openpyxl
from docx import Document

def wybierz_plik_excel():
    root = tk.Tk()
    root.withdraw()  # Ukryj okno główne
    sciezka = filedialog.askopenfilename(
        title="Wybierz plik Excel",
        filetypes=[("Pliki Excel", "*.xlsx *.xls")]
    )
    return sciezka

def odczytaj_dane_z_excela(plik_excel):
    wb = openpyxl.load_workbook(plik_excel)
    ws = wb.active

    nazwisko = ws['A2'].value
    wymiar = ws['B2'].value
    krawedz = ws['C2'].value
    kolor = ws['D2'].value
    nogi = ws['E2'].value
    terminOddania = ws['F2'].value

    dane = {
        'nazwisko': str(nazwisko),
        'wymiar': str(wymiar),
        'krawedz': str(krawedz),
        'kolor': str(kolor),
        'nogi': str(nogi),
        'terminOddania': str(terminOddania),
    }
    return dane

def zamien_teksty_w_runach(para, dane):
    for run in para.runs:
        for klucz, wartosc in dane.items():
            placeholder = f"{{{{{klucz}}}}}"
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, wartosc)

def wypelnij_word(dane, szablon_docx, wyjscie_docx):
    doc = Document(szablon_docx)
    for para in doc.paragraphs:
        zamien_teksty_w_runach(para, dane)
    for tabela in doc.tables:
        for wiersz in tabela.rows:
            for komorka in wiersz.cells:
                for para in komorka.paragraphs:
                    zamien_teksty_w_runach(para, dane)
    doc.save(wyjscie_docx)
    print(f"✅ Utworzono plik: {wyjscie_docx}")

if __name__ == "__main__":
    plik_excel = wybierz_plik_excel()
    if plik_excel:
        dane = odczytaj_dane_z_excela(plik_excel)
        wypelnij_word(dane, "szablon.docx", "wypelniony_raport.docx")
    else:
        print("Nie wybrano pliku.")
